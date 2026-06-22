# -*- coding: utf-8 -*-

import os
import sys

models_dir = os.path.dirname(os.path.abspath(__file__))
backend_dir = os.path.dirname(models_dir)
if backend_dir not in sys.path:
    sys.path.insert(0, backend_dir)

import logging
from typing import Optional
from docx import Document

logger = logging.getLogger(__name__)

class FileReader:
    """
    文档内容提取工具类
    """
    
    @staticmethod
    def extract_text(file_path: str) -> str:
        """
        根据文件后缀名，从文件中提取文本内容
        支持 .docx, .txt, .md
        :param file_path: 文件路径
        :return: 提取出的文本字符串
        """
        if not os.path.exists(file_path):
            logger.error(f"文件不存在: {file_path}")
            return ""

        ext = os.path.splitext(file_path)[1].lower()
        
        if ext in [".docx", ".doc"]:
            return FileReader._extract_docx(file_path)
        elif ext in [".txt", ".md"]:
            return FileReader._extract_plain_text(file_path)
        elif ext == ".pdf":
            return FileReader._extract_pdf(file_path)
        else:
            logger.error(f"不支持的文件格式: {ext}")
            return ""

    @staticmethod
    def _extract_docx(file_path: str) -> str:
        try:
            doc = Document(file_path)
            full_text = []
            
            # 1. 提取所有段落内容
            for para in doc.paragraphs:
                cleaned_text = para.text.strip()
                if cleaned_text:
                    full_text.append(cleaned_text)
                    
            # 2. 提取表格中的内容
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cleaned_cell = cell.text.strip()
                        if cleaned_cell:
                            full_text.append(cleaned_cell)
                            
            return "\n".join(full_text)
        except Exception as e:
            logger.error(f"解析 Word 文档失败: {e}")
            return ""

    @staticmethod
    def _extract_plain_text(file_path: str) -> str:
        try:
            # 尝试使用 multiple encodings
            for encoding in ['utf-8', 'gbk', 'utf-16']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            return ""
        except Exception as e:
            logger.error(f"解析文本文件失败: {e}")
            return ""

    @staticmethod
    def _extract_pdf(file_path: str) -> str:
        try:
            import PyPDF2
            text = ""
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            return text.strip()
        except ImportError:
            logger.error("未安装 PyPDF2，无法解析 PDF 文件。请运行: pip install PyPDF2")
            return "[错误: 后端未安装 PDF 解析插件]"
        except Exception as e:
            logger.error(f"解析 PDF 失败: {e}")
            return ""

    @staticmethod
    def format_as_prompt(filename: str, content: str) -> str:
        """
        将提取的内容格式化为 LLM Prompt 友好的字符串
        """
        if not content:
            return ""
            
        return (
            f"\n--- [上传的文件内容开始] ---\n"
            f"文件名: {filename}\n"
            f"内容如下:\n"
            f"{content}\n"
            f"--- [上传的文件内容结束] ---\n"
        )
